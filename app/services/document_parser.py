"""
Document parsing service for extracting text from PDF and DOC/DOCX files.

Uses pdfminer.six for PDF extraction (better layout preservation than PyPDF2)
and python-docx for DOCX with table extraction and header/footer stripping.
"""

import re
import unicodedata
from collections import Counter
from io import BytesIO
from typing import List, Optional

import logging

logger = logging.getLogger(__name__)

# Common ligature replacements
LIGATURES = {
    "\ufb01": "fi",
    "\ufb02": "fl",
    "\ufb00": "ff",
    "\ufb03": "ffi",
    "\ufb04": "ffl",
}

# Page number patterns to strip
PAGE_NUMBER_PATTERNS = [
    re.compile(r"^\s*-\s*\d+\s*-\s*$"),          # - 1 -
    re.compile(r"^\s*Page\s+\d+\s+of\s+\d+\s*$", re.IGNORECASE),  # Page 1 of 5
    re.compile(r"^\s*\d+\s*/\s*\d+\s*$"),         # 1/5
    re.compile(r"^\s*\d{1,3}\s*$"),               # standalone page number
]


class DocumentParser:
    """Service for parsing and extracting text from documents."""

    def extract_text(self, file_content: bytes, mime_type: str) -> Optional[str]:
        """
        Extract text from document content.

        Args:
            file_content: Document content as bytes
            mime_type: MIME type of the document

        Returns:
            Extracted text or None if extraction fails
        """
        try:
            if mime_type == "application/pdf":
                return self._extract_from_pdf(file_content)
            elif mime_type in [
                "application/msword",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ]:
                return self._extract_from_doc(file_content, mime_type)
            else:
                logger.warning(f"Unsupported MIME type for text extraction: {mime_type}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from document: {e}")
            return None

    def _extract_from_pdf(self, file_content: bytes) -> Optional[str]:
        """
        Extract text from PDF using pdfminer.six (with PyPDF2 fallback).

        Args:
            file_content: PDF content as bytes

        Returns:
            Extracted text or None if extraction fails
        """
        text = None

        # Primary: pdfminer.six — better layout and text extraction
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract

            text = pdfminer_extract(BytesIO(file_content))
        except ImportError:
            logger.info("pdfminer.six not installed, falling back to PyPDF2")
        except Exception as e:
            logger.warning(f"pdfminer extraction failed, trying PyPDF2: {e}")

        # Fallback: PyPDF2
        if not text:
            try:
                from PyPDF2 import PdfReader

                reader = PdfReader(BytesIO(file_content))
                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                text = "\n\n".join(text_parts)
            except ImportError:
                logger.warning("Neither pdfminer.six nor PyPDF2 available for PDF extraction")
                return None
            except Exception as e:
                logger.error(f"Error extracting text from PDF: {e}")
                return None

        if not text or not text.strip():
            return None

        text = self._strip_headers_footers(text)
        text = self._normalize_text(text)
        return text.strip() if text else None

    def _extract_from_doc(self, file_content: bytes, mime_type: str) -> Optional[str]:
        """
        Extract text from DOC/DOCX using python-docx (with table support).

        Args:
            file_content: Document content as bytes
            mime_type: MIME type (to differentiate DOC from DOCX)

        Returns:
            Extracted text or None if extraction fails
        """
        try:
            # DOCX extraction (modern format)
            if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                try:
                    from docx import Document

                    doc_file = BytesIO(file_content)
                    document = Document(doc_file)

                    text_parts = []

                    # Extract paragraph text
                    for paragraph in document.paragraphs:
                        if paragraph.text:
                            text_parts.append(paragraph.text)

                    # Extract table text
                    for table in document.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                if cell_text:
                                    row_text.append(cell_text)
                            if row_text:
                                text_parts.append(" | ".join(row_text))

                    # Strip header/footer content from DOCX sections
                    header_footer_text = set()
                    for section in document.sections:
                        if section.header and section.header.text:
                            header_footer_text.add(section.header.text.strip())
                        if section.footer and section.footer.text:
                            header_footer_text.add(section.footer.text.strip())

                    # Remove header/footer lines from extracted text
                    if header_footer_text:
                        text_parts = [
                            p for p in text_parts
                            if p.strip() not in header_footer_text
                        ]

                    extracted_text = "\n\n".join(text_parts)
                    extracted_text = self._normalize_text(extracted_text)
                    return extracted_text.strip() if extracted_text else None

                except ImportError:
                    logger.warning("python-docx not installed. DOCX text extraction disabled.")
                    return None

            # DOC extraction (legacy format)
            elif mime_type == "application/msword":
                try:
                    import textract

                    text = textract.process(BytesIO(file_content))
                    return text.decode('utf-8').strip() if text else None

                except ImportError:
                    logger.warning("textract not installed. DOC text extraction disabled.")
                    return None
                except Exception as e:
                    logger.error(f"Error extracting text from DOC: {e}")
                    return None

        except Exception as e:
            logger.error(f"Error extracting text from document: {e}")
            return None

    def _strip_headers_footers(self, text: str) -> str:
        """
        Remove repeated headers/footers and page numbers from extracted text.

        Splits text by form-feed characters (page boundaries), identifies lines
        that repeat across 3+ pages, and removes them.
        """
        pages = text.split("\f")
        if len(pages) < 2:
            # Single page — just strip page number patterns
            return self._strip_page_numbers(text)

        # Collect first and last 3 lines of each page
        candidate_lines: Counter = Counter()
        for page in pages:
            lines = [l.strip() for l in page.strip().split("\n") if l.strip()]
            if not lines:
                continue
            # Check first 3 and last 3 lines as header/footer candidates
            top = lines[:3]
            bottom = lines[-3:] if len(lines) > 3 else []
            for line in top + bottom:
                # Normalize whitespace for comparison
                normalized = re.sub(r"\s+", " ", line).strip()
                if normalized and len(normalized) > 1:
                    candidate_lines[normalized] += 1

        # Lines appearing on 3+ pages (or >50% of pages) are headers/footers
        threshold = min(3, max(2, len(pages) // 2))
        repeated = {line for line, count in candidate_lines.items() if count >= threshold}

        if not repeated:
            return self._strip_page_numbers(text)

        # Remove repeated lines from all pages
        cleaned_pages = []
        for page in pages:
            lines = page.split("\n")
            cleaned_lines = []
            for line in lines:
                normalized = re.sub(r"\s+", " ", line).strip()
                if normalized not in repeated:
                    cleaned_lines.append(line)
            cleaned_pages.append("\n".join(cleaned_lines))

        result = "\n\n".join(cleaned_pages)
        return self._strip_page_numbers(result)

    def _strip_page_numbers(self, text: str) -> str:
        """Remove standalone page number patterns from text."""
        lines = text.split("\n")
        cleaned = []
        for line in lines:
            is_page_num = any(p.match(line) for p in PAGE_NUMBER_PATTERNS)
            if not is_page_num:
                cleaned.append(line)
        return "\n".join(cleaned)

    def _normalize_text(self, text: str) -> str:
        """
        Normalize extracted text: fix ligatures, unicode, and layout issues.
        """
        # Fix common ligatures
        for lig, replacement in LIGATURES.items():
            text = text.replace(lig, replacement)

        # Unicode NFKD normalization (decomposes compatibility characters)
        text = unicodedata.normalize("NFKD", text)

        # Collapse multi-column layouts: short adjacent lines that look like columns
        lines = text.split("\n")
        merged: List[str] = []
        i = 0
        while i < len(lines):
            line = lines[i]
            # Detect potential column layout: very short lines followed by more short lines
            # with large gaps (multiple spaces in the middle)
            if "   " in line and len(line) > 20:
                # Possible multi-column line — split by large whitespace gaps
                parts = re.split(r"\s{3,}", line)
                if len(parts) >= 2 and all(len(p.strip()) < 60 for p in parts):
                    # Treat each part as a separate line
                    for part in parts:
                        part = part.strip()
                        if part:
                            merged.append(part)
                    i += 1
                    continue
            merged.append(line)
            i += 1

        text = "\n".join(merged)

        # Normalize whitespace
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"\r", "\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text

    def extract_metadata(self, file_content: bytes, mime_type: str) -> dict:
        """
        Extract metadata from document (title, author, etc.).

        Args:
            file_content: Document content as bytes
            mime_type: MIME type of the document

        Returns:
            Dictionary of metadata
        """
        metadata = {}

        try:
            if mime_type == "application/pdf":
                metadata = self._extract_pdf_metadata(file_content)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                metadata = self._extract_docx_metadata(file_content)
        except Exception as e:
            logger.error(f"Error extracting metadata: {e}")

        return metadata

    def _extract_pdf_metadata(self, file_content: bytes) -> dict:
        """Extract metadata from PDF."""
        try:
            from PyPDF2 import PdfReader

            pdf_file = BytesIO(file_content)
            reader = PdfReader(pdf_file)

            metadata = {}
            if reader.metadata:
                metadata = {
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'subject': reader.metadata.get('/Subject', ''),
                    'creator': reader.metadata.get('/Creator', ''),
                }

            metadata['num_pages'] = len(reader.pages)
            return metadata

        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
            return {}

    def _extract_docx_metadata(self, file_content: bytes) -> dict:
        """Extract metadata from DOCX."""
        try:
            from docx import Document

            doc_file = BytesIO(file_content)
            document = Document(doc_file)

            metadata = {}
            core_properties = document.core_properties

            metadata = {
                'title': core_properties.title or '',
                'author': core_properties.author or '',
                'subject': core_properties.subject or '',
                'keywords': core_properties.keywords or '',
            }

            return metadata

        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
            return {}


# Singleton instance
document_parser = DocumentParser()
